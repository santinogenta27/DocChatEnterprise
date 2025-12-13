"""
Multi-Strategy Parsing: Extracción de contenido usando múltiples estrategias
Basado en "Advanced ingestion process powered by LLM parsing for RAG system" y
"Hybrid OCR-LLM Framework for Enterprise-Scale Document Information Extraction"

Estrategias:
1. FAST: Python libraries (PyMuPDF, pdfplumber)
2. LLM: Multimodal LLM para OCR (Claude Sonnet, GPT-4V)
3. OCR: External OCR engines (PaddleOCR, EasyOCR, AWS Textract)
"""

from __future__ import annotations

import json
from typing import List, Dict, Optional, Any, Tuple
from pathlib import Path
from dataclasses import dataclass, field
from enum import Enum

from langchain_core.documents import Document


class ParsingStrategy(Enum):
    """Estrategias de parsing disponibles."""
    FAST = "fast"  # Python libraries
    LLM = "llm"  # Multimodal LLM
    OCR = "ocr"  # External OCR engines


class DocumentFormat(Enum):
    """Formatos de documento soportados."""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    PNG = "png"
    JPG = "jpg"
    MD = "md"


@dataclass
class ParsingResult:
    """Resultado de parsing de una estrategia."""
    strategy: ParsingStrategy
    text: str
    images: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    success: bool = True
    error: Optional[str] = None
    processing_time: float = 0.0


@dataclass
class Node:
    """Nodo base para estructura jerárquica de documentos."""
    node_id: str
    node_type: str  # Header, Text, Table, Image, Page, Document
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    children: List['Node'] = field(default_factory=list)
    parent: Optional['Node'] = None
    relationships: Dict[str, List[str]] = field(default_factory=dict)  # next, previous, parent, child


class MultiStrategyParser:
    """
    Parser que combina múltiples estrategias para extraer contenido de documentos.
    
    Estrategias:
    - FAST: Usa bibliotecas Python (PyMuPDF, pdfplumber, markitdown)
    - LLM: Usa modelos multimodales (Claude Sonnet, GPT-4V) para OCR y descripción
    - OCR: Usa motores OCR especializados (PaddleOCR, EasyOCR, AWS Textract)
    """
    
    def __init__(
        self,
        llm: Optional[Any] = None,
        use_fast: bool = True,
        use_llm: bool = True,
        use_ocr: bool = True
    ):
        self.llm = llm
        self.use_fast = use_fast
        self.use_llm = use_llm
        self.use_ocr = use_ocr
        
        # Inicializar parsers según disponibilidad
        self.fast_parser = None
        self.ocr_parser = None
        
        if use_fast:
            self._init_fast_parser()
        
        if use_ocr:
            self._init_ocr_parser()
    
    def _init_fast_parser(self):
        """Inicializa parsers rápidos (Python libraries)."""
        try:
            import fitz  # PyMuPDF
            self.fast_parser = "pymupdf"
        except ImportError:
            try:
                import pdfplumber
                self.fast_parser = "pdfplumber"
            except ImportError:
                self.fast_parser = None
    
    def _init_ocr_parser(self):
        """Inicializa parsers OCR."""
        try:
            from paddleocr import PaddleOCR
            self.ocr_parser = "paddleocr"
        except ImportError:
            try:
                import easyocr
                self.ocr_parser = "easyocr"
            except ImportError:
                self.ocr_parser = None
    
    def parse_document(
        self,
        file_path: str,
        doc_format: DocumentFormat
    ) -> Dict[ParsingStrategy, ParsingResult]:
        """
        Parsea un documento usando múltiples estrategias y combina resultados.
        
        Returns:
            Diccionario con resultados de cada estrategia
        """
        results = {}
        
        # 1. FAST parsing
        if self.use_fast and self.fast_parser:
            results[ParsingStrategy.FAST] = self._parse_fast(file_path, doc_format)
        
        # 2. LLM parsing (si hay LLM disponible)
        if self.use_llm and self.llm:
            results[ParsingStrategy.LLM] = self._parse_llm(file_path, doc_format)
        
        # 3. OCR parsing
        if self.use_ocr and self.ocr_parser:
            results[ParsingStrategy.OCR] = self._parse_ocr(file_path, doc_format)
        
        return results
    
    def _parse_fast(self, file_path: str, doc_format: DocumentFormat) -> ParsingResult:
        """Parsing rápido usando bibliotecas Python."""
        import time
        start_time = time.time()
        
        try:
            if doc_format == DocumentFormat.PDF:
                if self.fast_parser == "pymupdf":
                    import fitz
                    doc = fitz.open(file_path)
                    text_parts = []
                    images = []
                    tables = []
                    
                    for page_num, page in enumerate(doc):
                        text = page.get_text()
                        text_parts.append(text)
                        
                        # Extraer imágenes
                        image_list = page.get_images()
                        for img_index, img in enumerate(image_list):
                            images.append({
                                "page": page_num + 1,
                                "index": img_index,
                                "xref": img[0]
                            })
                    
                    return ParsingResult(
                        strategy=ParsingStrategy.FAST,
                        text="\n\n".join(text_parts),
                        images=images,
                        tables=tables,
                        processing_time=time.time() - start_time
                    )
                elif self.fast_parser == "pdfplumber":
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text_parts = []
                        tables = []
                        
                        for page in pdf.pages:
                            text_parts.append(page.extract_text() or "")
                            
                            # Extraer tablas
                            page_tables = page.extract_tables()
                            for table in page_tables:
                                tables.append({
                                    "data": table,
                                    "page": page.page_number
                                })
                        
                        return ParsingResult(
                            strategy=ParsingStrategy.FAST,
                            text="\n\n".join(text_parts),
                            tables=tables,
                            processing_time=time.time() - start_time
                        )
            
            elif doc_format == DocumentFormat.DOCX:
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text_parts = [para.text for para in doc.paragraphs]
                    
                    return ParsingResult(
                        strategy=ParsingStrategy.FAST,
                        text="\n".join(text_parts),
                        processing_time=time.time() - start_time
                    )
                except ImportError:
                    pass
            
            return ParsingResult(
                strategy=ParsingStrategy.FAST,
                text="",
                success=False,
                error=f"Fast parser not available for {doc_format.value}",
                processing_time=time.time() - start_time
            )
        
        except Exception as e:
            return ParsingResult(
                strategy=ParsingStrategy.FAST,
                text="",
                success=False,
                error=str(e),
                processing_time=time.time() - start_time
            )
    
    def _parse_llm(self, file_path: str, doc_format: DocumentFormat) -> ParsingResult:
        """Parsing usando LLM multimodal."""
        import time
        start_time = time.time()
        
        if not self.llm:
            return ParsingResult(
                strategy=ParsingStrategy.LLM,
                text="",
                success=False,
                error="LLM not available",
                processing_time=time.time() - start_time
            )
        
        try:
            # Para imágenes, usar visión del LLM
            if doc_format in [DocumentFormat.PNG, DocumentFormat.JPG]:
                # Leer imagen como base64
                import base64
                with open(file_path, "rb") as img_file:
                    img_data = base64.b64encode(img_file.read()).decode()
                
                # Prompt para LLM multimodal
                prompt = f"""Extract all text, tables, and describe images from this document image.
Return a structured JSON with:
- text: extracted text content
- tables: list of tables found
- images: descriptions of images found

Image data: {img_data[:100]}..."""
                
                response = self.llm.invoke(prompt)
                result_data = json.loads(str(response.content))
                
                return ParsingResult(
                    strategy=ParsingStrategy.LLM,
                    text=result_data.get("text", ""),
                    tables=result_data.get("tables", []),
                    images=result_data.get("images", []),
                    processing_time=time.time() - start_time
                )
            
            # Para PDFs, primero extraer con FAST y luego mejorar con LLM
            fast_result = self._parse_fast(file_path, doc_format)
            
            if fast_result.success:
                # Mejorar extracción con LLM
                prompt = f"""Improve and structure this extracted text from a PDF document.
Fix any OCR errors, organize by sections, and extract tables and images.

Text:
{fast_result.text[:2000]}"""
                
                response = self.llm.invoke(prompt)
                improved_text = str(response.content)
                
                return ParsingResult(
                    strategy=ParsingStrategy.LLM,
                    text=improved_text,
                    images=fast_result.images,
                    tables=fast_result.tables,
                    processing_time=time.time() - start_time
                )
            
            return ParsingResult(
                strategy=ParsingStrategy.LLM,
                text="",
                success=False,
                error="Failed to parse with FAST first",
                processing_time=time.time() - start_time
            )
        
        except Exception as e:
            return ParsingResult(
                strategy=ParsingStrategy.LLM,
                text="",
                success=False,
                error=str(e),
                processing_time=time.time() - start_time
            )
    
    def _parse_ocr(self, file_path: str, doc_format: DocumentFormat) -> ParsingResult:
        """Parsing usando OCR engines."""
        import time
        start_time = time.time()
        
        if not self.ocr_parser:
            return ParsingResult(
                strategy=ParsingStrategy.OCR,
                text="",
                success=False,
                error="OCR parser not available",
                processing_time=time.time() - start_time
            )
        
        try:
            if doc_format in [DocumentFormat.PNG, DocumentFormat.JPG]:
                if self.ocr_parser == "paddleocr":
                    from paddleocr import PaddleOCR
                    ocr = PaddleOCR(use_angle_cls=True, lang='en')
                    result = ocr.ocr(file_path, cls=True)
                    
                    text_parts = []
                    for line in result[0] if result else []:
                        if line and len(line) >= 2:
                            text_parts.append(line[1][0])
                    
                    return ParsingResult(
                        strategy=ParsingStrategy.OCR,
                        text="\n".join(text_parts),
                        processing_time=time.time() - start_time
                    )
                
                elif self.ocr_parser == "easyocr":
                    import easyocr
                    reader = easyocr.Reader(['en'])
                    result = reader.readtext(file_path)
                    
                    text_parts = [item[1] for item in result]
                    
                    return ParsingResult(
                        strategy=ParsingStrategy.OCR,
                        text="\n".join(text_parts),
                        processing_time=time.time() - start_time
                    )
            
            return ParsingResult(
                strategy=ParsingStrategy.OCR,
                text="",
                success=False,
                error=f"OCR not supported for {doc_format.value}",
                processing_time=time.time() - start_time
            )
        
        except Exception as e:
            return ParsingResult(
                strategy=ParsingStrategy.OCR,
                text="",
                success=False,
                error=str(e),
                processing_time=time.time() - start_time
            )
    
    def assemble_results(
        self,
        results: Dict[ParsingStrategy, ParsingResult]
    ) -> str:
        """
        Ensambla resultados de múltiples estrategias en un markdown unificado.
        Basado en "Multimodal Assembler Agent" del paper.
        """
        assembled_parts = []
        
        # Priorizar: LLM > OCR > FAST
        if ParsingStrategy.LLM in results and results[ParsingStrategy.LLM].success:
            llm_result = results[ParsingStrategy.LLM]
            assembled_parts.append(f"# LLM Extracted Content\n\n{llm_result.text}")
            
            # Agregar tablas
            if llm_result.tables:
                assembled_parts.append("\n## Tables\n\n")
                for table in llm_result.tables:
                    assembled_parts.append(f"```\n{table}\n```\n")
        
        elif ParsingStrategy.OCR in results and results[ParsingStrategy.OCR].success:
            ocr_result = results[ParsingStrategy.OCR]
            assembled_parts.append(f"# OCR Extracted Content\n\n{ocr_result.text}")
        
        elif ParsingStrategy.FAST in results and results[ParsingStrategy.FAST].success:
            fast_result = results[ParsingStrategy.FAST]
            assembled_parts.append(f"# Fast Parser Content\n\n{fast_result.text}")
            
            # Agregar tablas
            if fast_result.tables:
                assembled_parts.append("\n## Tables\n\n")
                for table in fast_result.tables:
                    assembled_parts.append(f"```\n{table}\n```\n")
        
        return "\n".join(assembled_parts)
    
    def create_node_structure(
        self,
        assembled_text: str,
        results: Dict[ParsingStrategy, ParsingResult]
    ) -> List[Node]:
        """
        Crea estructura de nodos jerárquica a partir del texto ensamblado.
        Basado en "Node-based extraction" del paper.
        """
        nodes = []
        
        # 1. Document Node (raíz)
        doc_node = Node(
            node_id="document_root",
            node_type="Document",
            content=assembled_text,
            metadata={"source": "multi_strategy_parsing"}
        )
        nodes.append(doc_node)
        
        # 2. Page Nodes
        pages = assembled_text.split("\n\n---\n\n")  # Separador de páginas
        for page_idx, page_content in enumerate(pages):
            page_node = Node(
                node_id=f"page_{page_idx + 1}",
                node_type="Page",
                content=page_content,
                parent=doc_node,
                metadata={"page_number": page_idx + 1}
            )
            nodes.append(page_node)
            doc_node.children.append(page_node)
        
        # 3. Header Nodes (detectar headers)
        import re
        header_pattern = r'^#+\s+(.+)$'
        current_page_node = None
        
        for line in assembled_text.split('\n'):
            header_match = re.match(header_pattern, line)
            if header_match:
                header_node = Node(
                    node_id=f"header_{len([n for n in nodes if n.node_type == 'Header'])}",
                    node_type="Header",
                    content=header_match.group(1),
                    parent=current_page_node or doc_node,
                    metadata={"level": len(line) - len(line.lstrip('#'))}
                )
                nodes.append(header_node)
                if current_page_node:
                    current_page_node.children.append(header_node)
        
        # 4. Table Nodes
        for strategy, result in results.items():
            for table_idx, table in enumerate(result.tables):
                table_node = Node(
                    node_id=f"table_{strategy.value}_{table_idx}",
                    node_type="Table",
                    content=json.dumps(table),
                    parent=doc_node,
                    metadata={"strategy": strategy.value, "index": table_idx}
                )
                nodes.append(table_node)
                doc_node.children.append(table_node)
        
        # 5. Image Nodes
        for strategy, result in results.items():
            for img_idx, image in enumerate(result.images):
                image_node = Node(
                    node_id=f"image_{strategy.value}_{img_idx}",
                    node_type="Image",
                    content=json.dumps(image),
                    parent=doc_node,
                    metadata={"strategy": strategy.value, "index": img_idx}
                )
                nodes.append(image_node)
                doc_node.children.append(image_node)
        
        # Establecer relaciones next/previous
        for i in range(len(nodes) - 1):
            nodes[i].relationships["next"] = [nodes[i + 1].node_id]
            nodes[i + 1].relationships["previous"] = [nodes[i].node_id]
        
        return nodes

